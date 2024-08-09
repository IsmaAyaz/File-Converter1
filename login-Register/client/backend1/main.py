# from fastapi import FastAPI, File, UploadFile, HTTPException
# from fastapi.responses import FileResponse
# import fitz  # PyMuPDF
# from docx import Document
# import os
# from fastapi.middleware.cors import CORSMiddleware
# import logging

# app = FastAPI()
# logging.basicConfig(
#     filename='app.log', 
#     level=logging.DEBUG, 
#     format='%(asctime)s - %(levelname)s - %(message)s'
# )


# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # Allows all origins; change this to specific origins in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# UPLOAD_DIR = "uploads"
# os.makedirs(UPLOAD_DIR, exist_ok=True)

# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     try:
#         file_location = os.path.join(UPLOAD_DIR, file.filename)
#         with open(file_location, "wb") as f:
#             f.write(await file.read())
#         return {"filename": file.filename, "filepath": file_location}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

# @app.post("/convert/")

# async def convert_file(filename: str):
#     file_path = os.path.join(UPLOAD_DIR, filename)
#     output_path = os.path.join(UPLOAD_DIR, f"converted_{filename}.txt")

#     try:
#         if file_path.lower().endswith('.pdf'):
#             convert_pdf_to_text(file_path, output_path)
#         elif file_path.lower().endswith('.docx'):
#             convert_word_to_text(file_path, output_path)  # Make sure to define this function
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format")

#         return {"output_path": output_path}
#     except HTTPException as e:
#         raise e
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error converting file: {str(e)}")

# @app.get("/download/")
# async def download_file(filepath: str):
#     if not os.path.exists(filepath):
#         raise HTTPException(status_code=404, detail="File not found")
#     return FileResponse(filepath, media_type='application/octet-stream', filename=os.path.basename(filepath))

# def convert_pdf_to_text(pdf_path, output_path):
#     try:
#         print(pdf_path)
#         pdf_document = fitz.open(pdf_path)
#         pdf_text = ""   
#         for page_num in range(len(pdf_document)):
#             page = pdf_document.load_page(page_num)
#             pdf_text += page.get_text("text")
#         with open(output_path, "w", encoding="utf-8") as text_file:
#             text_file.write(pdf_text)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error converting PDF to text: {str(e)}")

# def convert_word_to_text(word_path, output_path):
#     try:
#         doc = Document(word_path)
#         word_text = ""
#         for paragraph in doc.paragraphs:
#             word_text += paragraph.text + "\n"
#         with open(output_path, "w", encoding="utf-8") as text_file:
#             text_file.write(word_text)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error converting Word to text: {str(e)}")

from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import FileResponse
import fitz  # PyMuPDF
from docx import Document
import os
from fastapi.middleware.cors import CORSMiddleware
import logging

app = FastAPI()
logging.basicConfig(
    filename='app.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins; change this to specific origins in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as f:
            f.write(await file.read())
        return {"filename": file.filename, "filepath": file_location}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

@app.post("/convert/")
async def convert_file(request: Request):
    data = await request.json()
    filename = data.get("filename")
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    file_path = os.path.join(UPLOAD_DIR, os.path.basename(filename))
    output_path = os.path.join(UPLOAD_DIR, f"converted_{os.path.basename(filename)}.txt")

    try:
        if file_path.lower().endswith('.pdf'):
            convert_pdf_to_text(file_path, output_path)
        elif file_path.lower().endswith('.docx'):
            convert_word_to_text(file_path, output_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")

        return {"output_path": output_path}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting file: {str(e)}")

@app.get("/download/")
async def download_file(filepath: str):
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(filepath, media_type='application/octet-stream', filename=os.path.basename(filepath))

def convert_pdf_to_text(pdf_path, output_path):
    try:
        pdf_document = fitz.open(pdf_path)
        pdf_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            pdf_text += page.get_text("text")
        with open(output_path, "w", encoding="utf-8") as text_file:
            text_file.write(pdf_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting PDF to text: {str(e)}")

def convert_word_to_text(word_path, output_path):
    try:
        doc = Document(word_path)
        word_text = ""
        for paragraph in doc.paragraphs:
            word_text += paragraph.text + "\n"
        with open(output_path, "w", encoding="utf-8") as text_file:
            text_file.write(word_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting Word to text: {str(e)}")
